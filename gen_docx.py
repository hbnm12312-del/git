from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.8)
section.right_margin = Cm(2.8)

ea = qn("w:eastAsia")

def add_run(p, text, size=11.5, bold=False, color=RGBColor(0x2c,0x2c,0x2c)):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    r.font.name = "\u5fae\u8f6f\u96c5\u9ed1"
    r.element.rPr.rFonts.set(ea, "\u5fae\u8f6f\u96c5\u9ed1")
    return r

def new_p():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    return p

def sec_title(text):
    p = new_p()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, text, size=11, bold=True, color=RGBColor(0xb0,0x7d,0x4b))

def body(text):
    p = new_p()
    add_run(p, text)

def bullet(text):
    p = new_p()
    p.paragraph_format.left_indent = Cm(0.5)
    add_run(p, "\u00b7  " + text)

# === 标题 ===
p = doc.add_paragraph()
add_run(p, "\u623f\u7389\u6db5", size=28, bold=True, color=RGBColor(0x1a,0x1a,0x1a))
add_run(p, "    \u6750\u6599\u6210\u578b\u53ca\u63a7\u5236\u5de5\u7a0b \u00b7 \u5927\u4e8c", size=14, color=RGBColor(0x8a,0x8a,0x8a))

# 分割线
p = new_p()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(18)
add_run(p, "\u2501" * 40, size=6, color=RGBColor(0xd4,0xa3,0x73))

# === 简介 ===
sec_title("\u7b80  \u4ecb")
body("\u5b9d\u9e21\u6587\u7406\u5b66\u9662 \u673a\u68b0\u5de5\u7a0b\u5b66\u9662 \u00b7 \u6750\u6599\u6210\u578b\u53ca\u63a7\u5236\u5de5\u7a0b\u4e13\u4e1a \u00b7 2024\u7ea7\u672c\u79d1\u751f\u3002")
body("\u5e0c\u671b\u5728\u672c\u79d1\u9636\u6bb5\u8fdb\u5165\u8bfe\u9898\u7ec4\uff0c\u4ece\u57fa\u7840\u5b9e\u9a8c\u505a\u8d77\uff0c\u79ef\u7d2f\u79d1\u7814\u4e0e\u9879\u76ee\u7ecf\u9a8c\u3002")

# === 软件工具 ===
sec_title("\u8f6f\u4ef6\u5de5\u5177")
body("AutoCAD   UG / NX   SolidWorks   Origin   Python   Office \u4e09\u5957\u4ef6")

# === 个人特点 ===
sec_title("\u4e2a\u4eba\u7279\u70b9")
bullet("\u5b66\u4e60\u4e3b\u52a8\uff0c\u52a8\u624b\u80fd\u529b\u5f3a")
bullet("\u505a\u4e8b\u8e0f\u5b9e\uff0c\u6709\u8010\u5fc3")
bullet("\u8f6f\u4ef6\u4e0a\u624b\u5feb")
bullet("\u613f\u610f\u4ece\u57fa\u7840\u5b9e\u9a8c\u505a\u8d77")

# === 进组期望 ===
sec_title("\u8fdb\u7ec4\u671f\u671b")

tbl = doc.add_table(rows=3, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

goals = [
    "\u2460  \u53c2\u4e0e\u91d1\u76f8\u6280\u80fd\u7ade\u8d5b\uff0c\u5728\u5b9e\u6218\u4e2d\u63d0\u5347\u6750\u6599\u5206\u6790\u80fd\u529b",
    "\u2461  \u534f\u52a9\u8bfe\u9898\u7ec4\u5b9e\u9a8c\u5de5\u4f5c\uff0c\u4e89\u53d6\u53c2\u4e0e\u8bba\u6587\u53d1\u8868",
    "\u2462  \u79ef\u7d2f\u9879\u76ee\u7ecf\u9a8c\uff0c\u4e3a\u65e5\u540e\u8003\u7814\u5960\u5b9a\u57fa\u7840"
]

for i, goal in enumerate(goals):
    cell = tbl.cell(i, 0)
    cell.text = ""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # 左边框棕色
    tcBorders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:color"), "D4A373")
    tcBorders.append(left)
    for edge in ["top", "bottom", "right"]:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        tcBorders.append(el)
    tcPr.append(tcBorders)
    # 背景色
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "FAF8F5")
    shading.set(qn("w:val"), "clear")
    tcPr.append(shading)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, goal)

# 底部信息
p = new_p()
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(6)
add_run(p, "\u2501" * 40, size=6, color=RGBColor(0xdd,0xdd,0xdd))

p = new_p()
add_run(p, "\u260f  ", size=11, bold=True, color=RGBColor(0xb0,0x7d,0x4b))
add_run(p, "\u8bf7\u586b\u5199\u4f60\u7684\u624b\u673a\u53f7 / \u5fae\u4fe1", size=11, color=RGBColor(0x99,0x99,0x99))

output_path = os.path.expanduser("~/Desktop/\u623f\u7389\u6db5-\u4e2a\u4eba\u7b80\u4ecb.docx")
doc.save(output_path)
print("OK saved to: " + output_path)
